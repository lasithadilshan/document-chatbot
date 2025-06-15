import PyPDF2
from docx import Document
import streamlit as st
from typing import List
import re
import io

class DocumentProcessor:
    """Handles document parsing and text extraction with improved error handling"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_file) -> str:
        """Extract text from uploaded PDF file with better error handling"""
        try:
            # Reset file pointer
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(docx_file) -> str:
        """Extract text from uploaded DOCX file"""
        try:
            # Reset file pointer
            docx_file.seek(0)
            doc = Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(txt_file) -> str:
        """Extract text from uploaded TXT file with encoding detection"""
        try:
            # Reset file pointer
            txt_file.seek(0)
            # Try UTF-8 first, then fall back to other encodings
            try:
                text = txt_file.read().decode('utf-8')
            except UnicodeDecodeError:
                txt_file.seek(0)
                text = txt_file.read().decode('latin-1')
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading TXT: {str(e)}")
            return ""
    
    def process_uploaded_file(self, uploaded_file) -> str:
        """Process uploaded file based on its type"""
        if uploaded_file is None:
            return ""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(uploaded_file)
        elif file_extension == 'txt':
            return self.extract_text_from_txt(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.,!?;:()\-]', '', text)
        
        return text.strip()
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks with smart boundary detection"""
        if not text:
            return []
        
        # Clean the text first
        text = DocumentProcessor.clean_text(text)
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If we're not at the end of the text, try to break at a good boundary
            if end < len(text):
                # Look for sentence boundaries first
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for paragraph boundaries
                    paragraph_end = text.rfind('\n', start, end)
                    if paragraph_end > start + chunk_size // 2:
                        end = paragraph_end + 1
                    else:
                        # Look for word boundaries
                        word_end = text.rfind(' ', start, end)
                        if word_end > start + chunk_size // 2:
                            end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks